"""
update_download_docx_v3.py
Updates C:\\Users\\nicit\\Downloads\\WILD_MECHANICS_PRODUCTION_PIPELINE2.docx with:
1. Top Header Safe Zone & Spacing (MarginV=105 for Brand, MarginV=165 for Title).
2. Distinct Font Colors (Diamond White for Brand, Electric Yellow for Title).
3. Curiosity-Driven Publishing Titles Policy (High emotional stakes, no dry facts).
4. Subtitle Safe Zone for YouTube Shorts (MarginV=460 / Y~1460).
"""

import sys
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")

import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor

target_file = Path(r"C:\Users\nicit\Downloads\WILD_MECHANICS_PRODUCTION_PIPELINE2.docx")

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title & Subtitle
title_p = doc.add_paragraph()
title_run = title_p.add_run("Wild Mechanics — Master Production Pipeline Specification")
title_run.font.name = "Calibri"
title_run.font.size = Pt(22)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(20, 80, 160)

sub_p = doc.add_paragraph()
sub_run = sub_p.add_run("Official Channel Production Standard (Battle-Tested on Scarface Jaguar 2k+ views)")
sub_run.font.name = "Calibri"
sub_run.font.size = Pt(12)
sub_run.font.italic = True
sub_run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Calibri"
        r.font.bold = True
        if level == 1:
            r.font.size = Pt(15)
            r.font.color.rgb = RGBColor(20, 80, 160)
        elif level == 2:
            r.font.size = Pt(13)
            r.font.color.rgb = RGBColor(30, 60, 120)
        else:
            r.font.size = Pt(11.5)
            r.font.color.rgb = RGBColor(50, 50, 50)
    return h

def add_bullet(bold_prefix, text):
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(bold_prefix)
    r1.font.bold = True
    r1.font.name = "Calibri"
    r1.font.size = Pt(11)
    r2 = p.add_run(" " + text)
    r2.font.name = "Calibri"
    r2.font.size = Pt(11)

def add_callout(title, text):
    p = doc.add_paragraph()
    r1 = p.add_run(f"⚠️ {title}: ")
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(180, 40, 40)
    r1.font.name = "Calibri"
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.italic = True
    r2.font.name = "Calibri"
    r2.font.size = Pt(11)

# 1. AUDIO & FOOTAGE MANDATE
add_heading("1. Original Audio & Visual Transformation Mandate", level=1)
add_bullet("Original Voice & Sound Rule:", "The narration voiceover and background sound MUST be the original authentic documentary audio (original narrator + natural synchronized wildlife ambience). Do NOT replace the main story with synthetic robotic AI TTS.")
add_bullet("Audio Pitch & Frequency Modulation:", "Apply subtle pitch shifting (~ +/- 0.5 to 1.0 semitones) and acoustic filtering (asetrate, atempo, or EQ) to alter the audio fingerprint so it sounds distinct from the raw TV broadcast master while preserving natural vocal richness.")
add_bullet("Visual Transformation:", "The video must look visually distinct from real TV footage via: (1) 4:5 Ghost Blur framing, (2) subtle saturation/contrast punch, (3) on-screen comic action badges, (4) top header branding, and (5) word-level yellow karaoke subtitles.")

# 2. DURATION POLICY
add_heading("2. Strict Duration Policy & Minimum Thresholds", level=1)
add_callout(
    "STRICT MINIMUM DURATION RULE",
    "NO video produced for Wild Mechanics must EVER be less than 60.0 seconds total. Videos under 60 seconds harm algorithmic push, watch-time monetization eligibility, and audience retention."
)
add_heading("A. BBC Documentary Clips (Strict 61s – 63s Total Architecture)", level=2)
add_bullet("Main Story Arc:", "Must run for exactly 58.0s – 60.0s of continuous authentic documentary footage with pitch-altered original audio.")
add_bullet("Appended Outro CTA:", "+ 2.0s – 3.0s dedicated dynamic Outro CTA with custom AI voiceover tailored to the video's topic.")
add_bullet("Total Master Duration:", "Must be exactly 61.0s – 63.0s total master length (never under 60.0s).")
add_bullet("Content ID Protection Mechanism:", "Continuous raw BBC documentary audio strictly cuts off at <= 60.0s and transitions into fresh stock video + AI voiceover, breaking continuous audio fingerprinting.")

add_heading("B. Non-BBC Clips: Smithsonian / Love Nature / Discovery / Terra Mater (72s – 88s Total)", level=2)
add_bullet("Main Story Arc:", "70.0s – 85.0s to deliver the full narrative build-up, acoustic breakdown, and climactic strike.")
add_bullet("Appended Outro CTA:", "+ 2.0s – 3.0s dedicated dynamic Outro CTA tailored to the topic.")
add_bullet("Total Master Duration:", "72.0s – 88.0s total master length.")

add_heading("C. Smart Trimming (AI Sentence Boundaries & Duration Snapping)", level=2)
add_bullet("Sentence Boundary Detection:", "Uses Whisper word-level timestamps to detect natural punctuation periods (., !, ?) and speech pauses (>350ms).")
add_bullet("Duration Target Snapping:", "Smart Trimmer must snap to the nearest completed sentence boundary right at the target ceiling (58.0s – 60.0s for BBC / 70.0s – 85.0s for Non-BBC). It must NEVER truncate prematurely into a 40s or 50s clip.")

# 3. ON-SCREEN BRANDING & CURIOSITY TITLES
add_heading("3. On-Screen Branding & Curiosity-Driven Titles", level=1)
add_bullet("Header Positioning (Safe Zone):", "The top branding header must sit comfortably in the upper blur zone above the 4:5 video boundary (Y=285). Place WILD MECHANICS at Y=105 and Title at Y=165, maintaining an 80px buffer above the video boundary.")
add_bullet("Distinct Header Font Colors:", "WILD MECHANICS is styled in Diamond White (#FFFFFF / &H00FFFFFF&, Arial Bold). The Episode Title is styled in high-voltage Electric Yellow (#FFFF00 / &H0000FFFF&, Impact Bold).")
add_bullet("Curiosity-Driven Title Strategy:", "Titles must NEVER be dry educational facts. They must create intense curiosity, intrigue, or emotional stakes (e.g. 'WHY SALMON JUMP INTO A BEAR'S MOUTH 😱', 'THEY WAITED 10 MONTHS FOR THIS 1-SECOND STRIKE 💥', 'THE DEADLIEST 3 FEET IN THE RIVER 🐻').")

# 4. CANVAS & FRAMING
add_heading("4. Canvas & Framing (4:5 Ghost Blur)", level=1)
add_bullet("Foreground Viewbox:", "4:5 Aspect Ratio (1080x1350) keeping the animal protagonist centered in the primary viewport (Y=285 to Y=1635).")
add_bullet("Background:", "Ambient 1080x1920 blurred background (boxblur=30:5, brightness=-0.08, saturation=1.15) rendered from the active frame.")
add_bullet("Color & Contrast Punch:", "Subtle saturation boost (1.12x) and contrast punch (1.04x) to make wildlife visuals vibrant and immersive on mobile screens.")
add_bullet("100% Watermark Elimination:", "The 4:5 center zoom/crop (iw-1080)/2:(ih-1350)/2 automatically crops out all corner broadcaster watermarks (PBS, BBC, Smithsonian, NatGeo).")

# 5. SUBTITLE STYLING & YOUTUBE SHORTS SAFE ZONES
add_heading("5. Subtitle Styling & YouTube Shorts Safe Zones", level=1)
add_bullet("Format:", "Advanced ASS Subtitles with millisecond \\k karaoke timing tags.")
add_bullet("Vertical Safe Zone (MarginV=460):", "Subtitles must be anchored at Y ≈ 1460 (MarginV=460 from bottom). This places them in the lower third of the 4:5 video box while staying safely ABOVE all YouTube Shorts bottom overlay buttons (channel handle, subscribe, audio title at Y > 1550).")
add_bullet("Typography:", "Heavy Condensed Bold (Impact / Montserrat ExtraBold, FontSize=58).")
add_bullet("Active Word Highlight:", "Electric Yellow (#FFFF00 / &H0000FFFF&) on the currently spoken word, transitioning back to clean white.")
add_bullet("Outline & Legibility:", "Solid 4–5px black outline (OutlineColour=&H00000000) for maximum legibility over dynamic animal motion.")

# 6. DYNAMIC TOPIC-MATCHED CTA
add_heading("6. Dynamic Topic-Matched CTA (Reference: u0TVV-v1Skg)", level=1)
add_bullet("Duration:", "Dedicated 2.5s – 3.0s Outro Clip stitched at the end of the video.")
add_bullet("Topic-Matched Video Background:", "Sourced from 100% clean, watermark-free high-resolution stock video specifically matching the featured animal.")
add_bullet("Voiceover Script:", "Punchy closing biological superpower statement + channel follow invite: 'Nature is full of hidden biological superpowers, just like this one. Follow Wild Mechanics for more.'")
add_bullet("Visual Typography (u0TVV-v1Skg Style):", "Centered, bold kinetic word bursts in Electric Yellow (FOLLOW -> WILD MECHANICS -> FOR MORE.) without static boxed cards.")

# 7. QA CHECKLIST
add_heading("7. Production Quality Assurance (QA) Checklist", level=1)
add_bullet("Total Duration Check:", "Is the master video >= 61.0s? (If < 60.0s, the render fails QA and must be re-rendered).")
add_bullet("Header Breathing Room:", "Is the title cleanly positioned above the 4:5 boundary without touching the edge?")
add_bullet("Subtitle Safe Zone:", "Are subtitles positioned high enough (Y ≈ 1460) to avoid YouTube Shorts bottom UI?")
add_bullet("Curiosity Title Check:", "Does the title evoke strong curiosity rather than presenting a dry fact?")
add_bullet("Watermark Check:", "Are all broadcaster logos 100% cropped out in the 4:5 viewbox?")

doc.save(str(target_file))
print(f"🎉 Fully Updated Word Document: {target_file} ({target_file.stat().st_size / 1024:.1f} KB)")
