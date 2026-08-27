"""
generate_docx_pipeline_doc.py
Generates a beautifully styled Microsoft Word (.docx) document of the Wild Mechanics Production Pipeline Specification.
"""

import sys
from pathlib import Path

# UTF-8 stdout
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")

import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Page Margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title & Subtitle
title_p = doc.add_paragraph()
title_run = title_p.add_run("Wild Mechanics — Master Production Pipeline Specification")
title_run.font.name = "Calibri"
title_run.font.size = Pt(22)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(20, 80, 160)

sub_p = doc.add_paragraph()
sub_run = sub_p.add_run("Official Channel Production Standard (Battle-Tested on Scarface Jaguar 2k+ views)")
sub_run.font.name = "Calibri"
sub_run.font.size = Pt(12)
sub_run.font.italic = True
sub_run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Calibri"
        r.font.bold = True
        if level == 1:
            r.font.size = Pt(15)
            r.font.color.rgb = RGBColor(20, 80, 160)
        else:
            r.font.size = Pt(13)
            r.font.color.rgb = RGBColor(40, 40, 40)
    return h

def add_bullet(bold_prefix, text):
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(bold_prefix)
    r1.font.bold = True
    r1.font.name = "Calibri"
    r1.font.size = Pt(11)
    r2 = p.add_run(" " + text)
    r2.font.name = "Calibri"
    r2.font.size = Pt(11)

# 1. Duration Policy
add_heading("1. Duration Policy", level=1)
add_heading("A. BBC Documentary Clips", level=2)
add_bullet("Main Story:", "Exactly <= 57.5s – 60s (Cold Hook + Core Subject Hunt Arc).")
add_bullet("Appended CTA:", "+ 2.0s – 3.0s dedicated dynamic Outro CTA with custom AI voiceover tailored to the video topic.")
add_bullet("Content ID Protection:", "Cutting documentary audio at <= 58s breaks continuous audio fingerprinting and guarantees 100% global clearance with zero worldwide blocks.")

add_heading("B. Non-BBC Clips (Smithsonian / Love Nature / Discovery / Terra Mater)", level=2)
add_bullet("Main Story:", "> 60s (typically 70s – 85s) to deliver the full narrative build-up and climactic strike without rushing or cutting speech.")
add_bullet("Appended CTA:", "+ 2.0s – 3.0s dedicated dynamic Outro CTA.")

# 2. Canvas & Framing
add_heading("2. Canvas & Framing (4:5 Ghost Blur)", level=1)
add_bullet("Foreground Viewbox:", "4:5 Aspect Ratio (1080x1350) keeping the animal protagonist centered in the primary viewport.")
add_bullet("Background:", "Ambient 1080x1920 blurred background (boxblur=30:5, brightness=-0.08, saturation=1.15).")
add_bullet("Zero Letterbox:", "No plain black letterbox bars allowed.")
add_bullet("Zero Broadcaster Watermarks:", "The 4:5 center zoom/crop (iw-1080)/2:(ih-1350)/2 automatically eliminates all broadcaster corner watermarks (PBS, BBC, Smithsonian).")

# 3. Subtitle Styling
add_heading("3. Subtitle Styling (Word-Level Kinetic Karaoke)", level=1)
add_bullet("Format:", "Advanced ASS Subtitles with millisecond \\k karaoke timing tags.")
add_bullet("Typography:", "Heavy Condensed Bold (Impact / Montserrat, FontSize=60).")
add_bullet("Active Highlight:", "Electric Yellow (#FFFF00 / &H0000FFFF&) on the currently spoken word, transitioning back to clean white.")
add_bullet("Outline:", "Solid 4–5px black outline for maximum legibility over dynamic animal motion.")
add_bullet("Synchronization:", "Extracted via Whisper word-level timestamps, guaranteed 100% time-synced with voiceover.")

# 4. On-Screen Layers
add_heading("4. On-Screen Layers & Comic Book Action Effects", level=1)
add_bullet("Top Header Branding:", "WILD MECHANICS (White font, black outline, Y=180) + [SPECIES | EPISODE TITLE] (Electric Yellow, black outline, Y=240).")
add_bullet("Dynamic Action Badges:", "Pop-up badges and emojis timed to peak action moments to create a dynamic comic book style effect (e.g. THE AMBUSH, TARGET LOCKED, THE STRIKE, CAIMAN KILL).")

# 5. Audio Stack
add_heading("5. Audio Stack & Pitch Treatment", level=1)
add_bullet("Master Audio:", "Authentic high-bitrate documentary audio + synchronized natural ambient sounds (water splashes, wing flaps, breath).")
add_bullet("Pitch Modulation:", "Subtle frequency/pitch adjustment applied to raw narration to differentiate the audio fingerprint from broadcast masters while maintaining natural acoustics.")
add_bullet("Loudness Normalization:", "Normalized to -14 LUFS (EBU R128) with True Peak <= -1.5 dBTP.")

# 6. Dynamic Topic-Matched CTA
add_heading("6. Dynamic Topic-Matched CTA (Reference: u0TVV-v1Skg)", level=1)
add_bullet("Duration:", "Dedicated 2.5s – 3.0s Outro Clip stitched at the end of the video.")
add_bullet("Dynamic Topic Matching:", "Sourced from stock video (Pexels / Pixabay) specifically matching the featured animal/topic (e.g. snowy owl for Owl short, underwater hunt for Shark, savanna sprint for Cheetah).")
add_bullet("Voiceover Script:", "Punchy closing biological superpower statement + channel follow invite: 'Nature is full of hidden biological superpowers, just like this one. Follow Wild Mechanics for more.'")
add_bullet("Visual Typography (u0TVV-v1Skg Style):", "Centered, bold kinetic word bursts in Electric Yellow (FOLLOW -> WILD MECHANICS -> FOR MORE.) without static boxed cards.")

out_docx = Path(r"c:\Users\nicit\.gemini\antigravity\scratch\OpenMontage\docs\WILD_MECHANICS_PRODUCTION_PIPELINE.docx")
doc.save(str(out_docx))
print(f"🎉 Generated Word Document: {out_docx} ({out_docx.stat().st_size / 1024:.1f} KB)")
